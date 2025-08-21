from django.contrib import admin
from django.utils.html import format_html
from django.http import HttpResponse
from .models import Layanan, RekomendasiTeknis, Intake, StatusRekomendasiTeknis
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


# ===== EXPORT WORD =====
@admin.action(description="Export selected to Word")
def export_to_word(modeladmin, request, queryset):
    document = Document()
    document.add_heading('Data Rekomendasi Teknis SDA', 0)

    for obj in queryset:
        document.add_heading(f"Pemohon: {obj.nama_pemohon or 'Tanpa Nama'}", level=1)
        document.add_paragraph(f"Kontak: {obj.kontak_pemohon or '-'}")
        document.add_paragraph(f"Email: {obj.email_pemohon or '-'}")
        document.add_paragraph(f"Layanan: {obj.layanan.nama_layanan if obj.layanan else '-'}")
        document.add_paragraph(f"Nama Perusahaan: {obj.nama_perusahaan or '-'}")
        document.add_paragraph(f"Nama Direktur: {obj.nama_direktur or '-'}")
        document.add_paragraph(f"Alamat Perusahaan: {obj.alamat_perusahaan or '-'}")
        document.add_paragraph(f"Tanggal Permohonan: {obj.tanggal_permohonan}")

        dok_fields = [
            'dok_gambar_desain', 'dok_izin_lingkungan', 'dok_berita_acara', 'dok_jenis_prasarana',
            'dok_kepemilikan_lahan', 'dok_perizinan_usaha', 'dok_proposal_teknis',
            'dok_rencana_operasi', 'dok_surat_permohonan'
        ]
        document.add_paragraph("Dokumen:")
        for field in dok_fields:
            file_field = getattr(obj, field)
            if file_field:
                document.add_paragraph(f"- {field.replace('dok_', '').replace('_', ' ').title()}: {file_field.name}")

        if obj.intakes.exists():
            document.add_paragraph("Intake Terkait:")
            for intake in obj.intakes.all():
                document.add_paragraph(
                    f"- Sumber Air: {intake.sumber_air or '-'}, "
                    f"Kelurahan/Desa: {intake.kelurahan_desa or '-'}, "
                    f"Kecamatan: {intake.kecamatan or '-'}, "
                    f"Volume: {intake.volume_pengambilan or '-'}"
                )

        document.add_paragraph("-" * 50)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="rekomtek.docx"'
    document.save(response)
    return response

# ===== EXPORT PDF (ReportLab) =====
@admin.action(description="Export selected to PDF")
def export_to_pdf(modeladmin, request, queryset):
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="rekomtek.pdf"'

    styles = getSampleStyleSheet()
    story = []

    for obj in queryset:
        # Header data pemohon
        story.append(Paragraph(f"<b>Pemohon:</b> {obj.nama_pemohon or 'Tanpa Nama'}", styles["Normal"]))
        story.append(Paragraph(f"<b>Kontak:</b> {obj.kontak_pemohon or '-'}", styles["Normal"]))
        story.append(Paragraph(f"<b>Email:</b> {obj.email_pemohon or '-'}", styles["Normal"]))
        story.append(Paragraph(f"<b>Layanan:</b> {obj.layanan.nama_layanan if obj.layanan else '-'}", styles["Normal"]))
        story.append(Paragraph(f"<b>Perusahaan:</b> {obj.nama_perusahaan or '-'}", styles["Normal"]))
        story.append(Paragraph(f"<b>Direktur:</b> {obj.nama_direktur or '-'}", styles["Normal"]))
        story.append(Paragraph(f"<b>Alamat:</b> {obj.alamat_perusahaan or '-'}", styles["Normal"]))
        story.append(Paragraph(f"<b>Tanggal Permohonan:</b> {obj.tanggal_permohonan}", styles["Normal"]))
        story.append(Spacer(1, 10))

        # Dokumen
        dok_fields = [
            'dok_gambar_desain', 'dok_izin_lingkungan', 'dok_berita_acara', 'dok_jenis_prasarana',
            'dok_kepemilikan_lahan', 'dok_perizinan_usaha', 'dok_proposal_teknis',
            'dok_rencana_operasi', 'dok_surat_permohonan'
        ]
        story.append(Paragraph("<b>Dokumen:</b>", styles["Normal"]))
        for field in dok_fields:
            file_field = getattr(obj, field)
            if file_field:
                file_url = request.build_absolute_uri(file_field.url)
                story.append(Paragraph(f"- {field.replace('dok_', '').replace('_', ' ').title()}: <link href='{file_url}'>{file_url}</link>", styles["Normal"]))
        story.append(Spacer(1, 10))

        # Intake terkait
        if obj.intakes.exists():
            story.append(Paragraph("<b>Intake Terkait:</b>", styles["Normal"]))
            for intake in obj.intakes.all():
                story.append(Paragraph(
                    f"- Sumber Air: {intake.sumber_air or '-'}, "
                    f"Kel/Desa: {intake.kelurahan_desa or '-'}, "
                    f"Kecamatan: {intake.kecamatan or '-'}, "
                    f"Volume: {intake.volume_pengambilan or '-'}",
                    styles["Normal"]
                ))
        story.append(Spacer(1, 15))
        story.append(Paragraph("<hr/>", styles["Normal"]))
        story.append(Spacer(1, 15))

    doc = SimpleDocTemplate(response, pagesize=A4)
    doc.build(story)
    return response

# ===== Inline Intake untuk RekomendasiTeknis =====
class IntakeInline(admin.TabularInline):
    model = Intake
    extra = 1

# ===== Admin RekomendasiTeknis =====
@admin.register(RekomendasiTeknis)
class RekomendasiTeknisAdmin(admin.ModelAdmin):
    list_display = ['nama_perusahaan', 'layanan', 'nama_pemohon', 'tanggal_permohonan', 'aksi']
    list_filter = ['layanan', 'tanggal_permohonan']
    search_fields = ['nama_perusahaan', 'nama_pemohon']
    inlines = [IntakeInline]
    actions = [export_to_word, export_to_pdf]

    def aksi(self, obj):
        return format_html(
            '<div style="display: flex; gap: 5px;">'
            '<a class="btn btn-warning btn-sm" href="{}">Edit</a>'
            '<a class="btn btn-danger btn-sm" href="{}">Hapus</a>'
            '</div>',
            f'/admin/rekomtek/rekomendasiteknis/{obj.id}/change/',
            f'/admin/rekomtek/rekomendasiteknis/{obj.id}/delete/',
        )
    aksi.short_description = 'Aksi'

# ===== Admin Lainnya =====
@admin.register(Intake)
class IntakeAdmin(admin.ModelAdmin):
    list_display = ['rekomtek', 'sumber_air', 'kelurahan_desa', 'kapasitas_pompa']
    list_filter = ['sumber_air', 'provinsi']
    search_fields = ['kelurahan_desa', 'kecamatan', 'kabupaten_kota']

@admin.register(Layanan)
class LayananAdmin(admin.ModelAdmin):
    list_display = ['nama_layanan']
    search_fields = ['nama_layanan']

@admin.register(StatusRekomendasiTeknis)
class StatusRekomendasiTeknisAdmin(admin.ModelAdmin):
    list_display = ['rekomtek', 'jadwal_kunjungan_lapangan', 'tanggal_kunjungan_lapangan_1', 'tanggal_kunjungan_lapangan_2', 'tanggal_kirim']
    list_filter = ['jadwal_kunjungan_lapangan']
    search_fields = ['rekomtek__nama_pemohon', 'rekomtek__nama_perusahaan']
